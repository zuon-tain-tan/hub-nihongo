import os
import json
import re  
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- HÀM HỖ TRỢ TẠO SỐ TRANG ---
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    """Hàm can thiệp vào XML của Word để tự động đánh số trang"""
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# --- HÀM CHÍNH ---
def tao_tai_lieu_mau():
    # 1. Khởi tạo tài liệu
    doc = Document()

    # 2. Thiết lập trang in (Page Setup) - Khổ A5 và Margin
    section = doc.sections[0]
    section.page_width = Cm(14.8)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.5)
    section.gutter = Cm(1.0) # Gutter 1 cm

    # 3. Cấu hình Font và Paragraph Mặc định
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.size = Pt(12)
    font.name = 'Times New Roman'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    
    p_format = style_normal.paragraph_format
    p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    # 4. Cấu hình Style Bảng (Line 1.0, Before/After 6pt)
    style_table = doc.styles.add_style('Bang_Style', 1)
    table_font = style_table.font
    table_font.name = 'Times New Roman'
    table_font.size = Pt(12)
    table_font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    table_p_format = style_table.paragraph_format
    table_p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    table_p_format.space_before = Pt(6)
    table_p_format.space_after = Pt(6)

    # 5. Cấu hình Heading 1, 2, 3 (Màu đen, Size 12, Line 1.5, Trước/Sau 0pt)
    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.color.rgb = RGBColor(0, 0, 0)
        h_font.name = 'Times New Roman'
        h_font.size = Pt(12)
        
        # FIX LỖI KẸT FONT CALIBRI BẰNG CÁCH XOÁ THEME FONT
        rFonts = h_font.element.rPr.rFonts
        if rFonts is not None:
            for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme']:
                if qn(theme_attr) in rFonts.attrib:
                    del rFonts.attrib[qn(theme_attr)]
            rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

        h_format = h_style.paragraph_format
        h_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h_format.space_before = Pt(0)
        h_format.space_after = Pt(0)
        
        # Heading 1 căn giữa
        if level == 1:
            h_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 6. Cấu hình đánh số trang ở Footer (Góc dưới bên phải)
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer_para.add_run())

    # ================= THÊM NỘI DUNG =================
    # TRANG 1: MỤC LỤC
    doc.add_heading('MỤC LỤC', level=1)
    doc.add_page_break() # Ngắt trang

    # TRANG 2: TỪ VỰNG N5
    doc.add_heading('TỪ VỰNG N5', level=1)

    # --- ĐỌC TOÀN BỘ FILE TRONG THƯ MỤC TUVUNG ---
    tuvung_dir = os.path.join(os.getcwd(), 'tuvung')
    
    if os.path.exists(tuvung_dir):
        # Lấy danh sách các file .json
        json_files = [f for f in os.listdir(tuvung_dir) if f.endswith('.json')]
        
        def get_number_prefix(filename):
            match = re.search(r'^(\d+)', filename)
            return int(match.group(1)) if match else 9999
        
        json_files.sort(key=get_number_prefix)
        
        # Duyệt qua từng file json (Có dùng enumerate để lấy vị trí index)
        for index, file_name in enumerate(json_files):
            json_path = os.path.join(tuvung_dir, file_name)
            
            h2_text = file_name.replace('.json', '')
            
            # Khởi tạo paragraph Heading 2
            h2_para = doc.add_heading(h2_text, level=2)
            
            # --- FIX LỖI "DÍNH Ở HEADING 2" ---
            # Nếu không phải là file JSON đầu tiên (tức là từ file thứ 2 trở đi),
            # tự động ép Heading 2 phải qua trang mới.
            if index > 0:
                h2_para.paragraph_format.page_break_before = True
            
            # Lấy ra chỉ mục của Heading 2 hiện tại (để nối vào Heading 3)
            chi_muc_h2 = get_number_prefix(file_name)

            # Đọc dữ liệu
            with open(json_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Nhóm dữ liệu theo 'phan_nhom'
            nhom_tu_vung = {}
            for item in data:
                nhom = item.get('phan_nhom', 'Khác')
                if nhom not in nhom_tu_vung:
                    nhom_tu_vung[nhom] = []
                nhom_tu_vung[nhom].append(item)
            
            # Biến đếm cho Heading 3
            chi_muc_h3 = 1
            
            # Render bảng cho từng nhóm trong file
            for nhom, danh_sach in nhom_tu_vung.items():
                
                # Khởi tạo paragraph Heading 3
                h3_para = doc.add_heading(f"{chi_muc_h2}.{chi_muc_h3}. {nhom}", level=3)
                
                # Ngắt trang cho Heading 3 (áp dụng từ mục .2 trở đi)
                if chi_muc_h3 > 1:
                    h3_para.paragraph_format.page_break_before = True
                    
                chi_muc_h3 += 1 
                
                # Tạo bảng 2 cột
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                
                # Tiêu đề cột
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Từ vựng'
                hdr_cells[1].text = 'Ý nghĩa'
                
                # Định dạng tiêu đề cột
                for cell in hdr_cells:
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
                    for p in cell.paragraphs:
                        p.style = 'Bang_Style'
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
                        for run in p.runs:
                            run.font.bold = True
                
                # Điền dữ liệu vào bảng
                for item in danh_sach:
                    row_cells = table.add_row().cells
                    row_cells[0].text = item.get('tu_vung', '')
                    row_cells[1].text = item.get('y_nghia', '')
                    
                    # Định dạng dữ liệu
                    for cell in row_cells:
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER 
                        for p in cell.paragraphs:
                            p.style = 'Bang_Style'
                            p.alignment = WD_ALIGN_PARAGRAPH.LEFT 
    else:
        doc.add_paragraph(f"Lỗi: Không tìm thấy thư mục: {tuvung_dir}")

    # 7. Lưu file
    file_path = os.path.join(os.getcwd(), 'TaiLieu_A5_Chuan.docx')
    doc.save(file_path)
    print(f"Đã tạo thành công file Word tại: {file_path}")

if __name__ == "__main__":
    tao_tai_lieu_mau()
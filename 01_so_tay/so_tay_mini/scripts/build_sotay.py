import os
import sys
import json
import re  
import logging
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT_ROOT = Path(__file__).resolve().parents[1]

if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')
if hasattr(sys.stdin, 'reconfigure'):
    sys.stdin.reconfigure(encoding='utf-8')

# --- CÁC HÀM HỖ TRỢ ---
def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'separate')
    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

def add_simple_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = instruction

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = 'Bấm Ctrl+A rồi F9 để cập nhật mục lục nếu cần.'

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.extend([fld_char_begin, instr_text, fld_char_separate, placeholder, fld_char_end])

def get_group_and_sort_key(text):
    """Bóc tách để tìm chữ cái đại diện cho Tiếng Nhật (Mềm -> Cứng -> Khác)"""
    if not text: return '*', '4_'
    t = text
    for char in ['[', ']', '~', '～', '(', ')', '「', '」', ' ', '…']:
        t = t.replace(char, '')
    t = t.strip()
    if not t: return '*', '4_'
        
    first_char = t[0]
    if '\u3040' <= first_char <= '\u309f': cat = '1_' 
    elif '\u30a0' <= first_char <= '\u30ff': cat = '2_' 
    else: cat = '3_' 
    
    # --- CẬP NHẬT: Đệm số 0 để sort tự nhiên (Natural Sorting) ---
    # 1 -> 00001, 10 -> 00010
    sort_str = re.sub(r'\d+', lambda m: m.group(0).zfill(5), t.lower())
    
    return first_char.upper(), cat + sort_str

def get_vietnamese_sort_key(text):
    """Bóc tách và khử dấu để tra từ theo Tiếng Việt (A-Z)"""
    if not text: return '#', 'zzz'
    
    t = text
    for char in ['[', ']', '~', '～', '(', ')', '「', '」', ' ', '…', '-', '"', "'"]:
        t = t.replace(char, '')
    t = t.strip()
    if not t: return '#', 'zzz'

    # Khử dấu Tiếng Việt
    s1 = u'ÀÁÂÃÈÉÊÌÍÒÓÔÕÙÚÝàáâãèéêìíòóôõùúýĂăĐđĨĩŨũƠơƯưẠạẢảẤấẦầẨẩẪẫẬậẮắẰằẲẳẴẵẶặẸẹẺẻẼẽẾếỀềỂểỄễỆệỈỉỊịỌọỎỏỐốỒồỔổỖỗỘộỚớỜờỞởỠỡỢợỤụỦủỨứỪừỬửỮữỰựỲỳỴỵỶỷỸỹ'
    s0 = u'AAAAEEEIIOOOOUUYaaaaeeeiioooouuyAaDdIiUuOoUuAaAaAaAaAaAaAaAaAaAaAaAaEeEeEeEeEeEeEeEeIiIiOoOoOoOoOoOoOoOoOoOoOoOoUuUuUuUuUuUuUuUuYyYyYyYy'
    
    t_clean = ''
    for c in t:
        if c in s1:
            t_clean += s0[s1.index(c)]
        else:
            t_clean += c
            
    first_char = t_clean[0].upper()
    
    if not ('A' <= first_char <= 'Z'):
        first_char = '#'
        
    # --- CẬP NHẬT: Đệm số 0 để sort tự nhiên (Natural Sorting) ---
    # VD: 1cai -> 00001cai , 10cai -> 00010cai
    sort_str = re.sub(r'\d+', lambda m: m.group(0).zfill(5), t_clean.lower())

    return first_char, sort_str

def get_number_prefix(filename):
    match = re.search(r'^(\d+)', filename)
    return int(match.group(1)) if match else 9999

def ask_export_options():
    valid_options = {'1', '2', '3'}
    print("Chọn phần muốn tạo trong sổ tay:")
    print("1. A-z chữ cái Hira,Kata")
    print("2. A-z chữ cái Vie")
    print("3. Gom cụm từ vựng")
    print("Ví dụ nhập: 1 hoặc 1,3 hoặc 1 2 3 hoặc all")

    while True:
        raw = input("Nhập lựa chọn: ").strip().lower()
        if raw in {'all', 'a', 'tat ca', 'tất cả'}:
            return {'1', '2', '3'}

        choices = {part for part in re.split(r'[\s,;]+', raw) if part}
        if choices and choices.issubset(valid_options):
            return choices

        print("Lựa chọn không hợp lệ. Ví dụ đúng: 1,2 hoặc 1 3 hoặc all.")

def apply_run_font(run, bold=False, size=12):
    run.font.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(14.8)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.0)
    section.right_margin = Cm(1.5)
    section.gutter = Cm(1.0)

    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.size = Pt(12)
    font.name = 'Times New Roman'
    font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

    p_format = style_normal.paragraph_format
    p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p_format.space_before = Pt(0)
    p_format.space_after = Pt(0)

    style_table = doc.styles.add_style('Bang_Style', 1)
    table_font = style_table.font
    table_font.name = 'Times New Roman'
    table_font.size = Pt(12)
    table_font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    table_p_format = style_table.paragraph_format
    table_p_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    table_p_format.space_before = Pt(6)
    table_p_format.space_after = Pt(6)

    style_index = doc.styles.add_style('Index_Style', 1)
    index_font = style_index.font
    index_font.name = 'Times New Roman'
    index_font.size = Pt(12)
    index_font.element.rPr.rFonts.set(qn('w:eastAsia'), 'Yu Gothic')
    index_p_format = style_index.paragraph_format
    index_p_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    index_p_format.space_before = Pt(0)
    index_p_format.space_after = Pt(0)
    tab_stops = index_p_format.tab_stops
    tab_stops.add_tab_stop(Cm(11.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

    for level in range(1, 4):
        h_style = doc.styles[f'Heading {level}']
        h_font = h_style.font
        h_font.color.rgb = RGBColor(0, 0, 0)
        h_font.name = 'Times New Roman'
        h_font.size = Pt(12)

        rFonts = h_font.element.rPr.rFonts
        if rFonts is not None:
            for theme_attr in ['w:asciiTheme', 'w:hAnsiTheme', 'w:cstheme']:
                if qn(theme_attr) in rFonts.attrib:
                    del rFonts.attrib[qn(theme_attr)]
            rFonts.set(qn('w:eastAsia'), 'Yu Gothic')

        h_format = h_style.paragraph_format
        h_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        h_format.space_before = Pt(0)
        h_format.space_after = Pt(0)
        if level == 1:
            h_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(footer_para.add_run())
    return doc

def add_toc_section(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(0)
    apply_run_font(title.add_run('MỤC LỤC'), bold=True)

    toc_para = doc.add_paragraph()
    add_simple_field(toc_para, r'TOC \o "1-3" \h \z \u')

def add_page_break_if_needed(doc, has_content):
    if has_content:
        doc.add_page_break()
    return True

def load_vocabulary_data(tuvung_dir):
    n5_dir = os.path.join(tuvung_dir, 'tu_vungN5')
    all_words = []
    json_data_map = []

    if not os.path.exists(n5_dir):
        return all_words, json_data_map

    json_files = [f for f in os.listdir(n5_dir) if f.endswith('.json')]
    json_files.sort(key=get_number_prefix)

    for file_name in json_files:
        json_path = os.path.join(n5_dir, file_name)
        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        for item in data:
            item_copy = item.copy()
            grp_jp, skey_jp = get_group_and_sort_key(item.get('tu_vung', ''))
            item_copy['jp_group_char'] = grp_jp
            item_copy['jp_sort_key'] = skey_jp

            grp_vn, skey_vn = get_vietnamese_sort_key(item.get('y_nghia', ''))
            item_copy['vn_group_char'] = grp_vn
            item_copy['vn_sort_key'] = skey_vn
            all_words.append(item_copy)

        json_data_map.append({
            'file_name': file_name,
            'data': data,
            'prefix_num': get_number_prefix(file_name)
        })

    return all_words, json_data_map

def load_5s_data(tuvung_dir):
    folder = os.path.join(tuvung_dir, '5S')
    if not os.path.exists(folder):
        return []

    data = []
    json_files = [f for f in os.listdir(folder) if f.endswith('.json')]
    json_files.sort(key=get_number_prefix)
    for file_name in json_files:
        json_path = os.path.join(folder, file_name)
        with open(json_path, 'r', encoding='utf-8') as f:
            data.extend(json.load(f))
    return data

def set_table_borders(table, color):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is not None:
        tbl_pr.remove(borders)

    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'single')
        tag.set(qn('w:sz'), '4')
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), color)
        borders.append(tag)
    tbl_pr.append(borders)

def set_cell_width(cell, width):
    cell.width = width
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in('w:tcW')
    if tc_w is None:
        tc_w = OxmlElement('w:tcW')
        tc_pr.append(tc_w)
    tc_w.set(qn('w:w'), str(int(width.cm * 567)))
    tc_w.set(qn('w:type'), 'dxa')

def add_table(doc, headers, rows, border_color=None, column_widths=None):
    header_count = 1 if headers else 0
    column_count = len(headers) if headers else len(rows[0]) if rows else 1
    table = doc.add_table(rows=header_count, cols=column_count)
    table.style = 'Table Grid'
    table.autofit = False

    if border_color:
        set_table_borders(table, border_color)

    if headers:
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = header
            if column_widths:
                set_cell_width(hdr_cells[idx], column_widths[idx])

        for cell in hdr_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.style = 'Bang_Style'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.bold = True

    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = value
            if column_widths:
                set_cell_width(row_cells[idx], column_widths[idx])

        for cell in row_cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.style = 'Bang_Style'
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_5s_paragraphs(doc, items):
    for index, item in enumerate(items):
        if index > 0:
            doc.add_paragraph()

        jp_para = doc.add_paragraph(style='Bang_Style')
        jp_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        jp_run = jp_para.add_run(item.get('tu_vung', ''))
        apply_run_font(jp_run)

        vn_para = doc.add_paragraph(style='Bang_Style')
        vn_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        vn_run = vn_para.add_run(item.get('y_nghia', ''))
        apply_run_font(vn_run)

def add_online_version_page(doc, qr_path, updated_date):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(24)
    title.paragraph_format.space_after = Pt(12)
    apply_run_font(title.add_run('Phiên bản Online của tài liệu:'), bold=True, size=14)

    qr_para = doc.add_paragraph()
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if qr_path.exists():
        qr_para.add_run().add_picture(str(qr_path), width=Cm(5), height=Cm(5))
    else:
        apply_run_font(qr_para.add_run('[Thiếu ảnh QR]'), bold=True)

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Pt(8)
    date_para.paragraph_format.space_after = Pt(42)
    apply_run_font(date_para.add_run(f'Cập nhật lần cuối: {updated_date}'), size=12)

    for _ in range(5):
        doc.add_paragraph()

    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.paragraph_format.space_after = Pt(6)
    apply_run_font(author_para.add_run('著者（ちょしゃ）： ズオン・タイン・タン'), bold=True, size=12)

    editor_para = doc.add_paragraph()
    editor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    apply_run_font(editor_para.add_run('編集協力（へんしゅうきょうりょく）： ChatGPT 5.5'), size=12)

def add_japanese_index(doc, all_words):
    doc.add_heading('TRA NHANH (HIRA, KATA)', level=1)
    current_group = ''
    for item in sorted(all_words, key=lambda x: x['jp_sort_key']):
        group_char = item['jp_group_char']
        if group_char != current_group:
            current_group = group_char
            p_char = doc.add_paragraph()
            apply_run_font(p_char.add_run(current_group), bold=True)
            p_char.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p_char.paragraph_format.space_before = Pt(0)
            p_char.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(style='Index_Style')
        p.text = f"{item.get('tu_vung', '')}\t{item.get('y_nghia', '')}"

def add_vietnamese_index(doc, all_words):
    doc.add_heading('TRA NHANH (TIẾNG VIỆT)', level=1)
    current_group = ''
    for item in sorted(all_words, key=lambda x: x['vn_sort_key']):
        group_char = item['vn_group_char']
        if group_char != current_group:
            current_group = group_char
            p_char = doc.add_paragraph()
            apply_run_font(p_char.add_run(current_group), bold=True)
            p_char.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            p_char.paragraph_format.space_before = Pt(0)
            p_char.paragraph_format.space_after = Pt(0)

        p = doc.add_paragraph(style='Index_Style')
        p.text = f"{item.get('y_nghia', '')}\t{item.get('tu_vung', '')}"

def add_5s_section(doc, data):
    if not data:
        return

    doc.add_heading('5S - 5K - 5C', level=1)
    grouped = {}
    for item in data:
        grouped.setdefault(item.get('phan_nhom', 'Khác'), []).append(item)

    for index, (group_name, items) in enumerate(grouped.items(), start=1):
        h2_para = doc.add_heading(group_name, level=2)
        if index > 1:
            h2_para.paragraph_format.page_break_before = True

        if group_name == '5S':
            add_5s_paragraphs(doc, items)
        else:
            rows = [(item.get('tu_vung', ''), item.get('y_nghia', '')) for item in items]
            add_table(doc, None, rows, border_color='FFFFFF', column_widths=[Cm(6.19), Cm(5.49)])

def add_grouped_vocabulary(doc, json_data_map):
    doc.add_heading('TỪ VỰNG N5 CHI TIẾT', level=1)
    for index, file_data in enumerate(json_data_map):
        h2_text = file_data['file_name'].replace('.json', '')
        chi_muc_h2 = file_data['prefix_num']

        h2_para = doc.add_heading(h2_text, level=2)
        if index > 0:
            h2_para.paragraph_format.page_break_before = True

        nhom_tu_vung = {}
        for item in file_data['data']:
            nhom = item.get('phan_nhom', 'Khác')
            nhom_tu_vung.setdefault(nhom, []).append(item)

        chi_muc_h3 = 1
        for nhom, danh_sach in nhom_tu_vung.items():
            h3_para = doc.add_heading(f"{chi_muc_h2}.{chi_muc_h3}. {nhom}", level=3)
            if chi_muc_h3 > 1:
                h3_para.paragraph_format.page_break_before = True

            chi_muc_h3 += 1
            rows = [(item.get('tu_vung', ''), item.get('y_nghia', '')) for item in danh_sach]
            add_table(doc, ['Từ vựng', 'Ý nghĩa'], rows)

def update_docx_fields_and_export_pdf(docx_path, pdf_path):
    import win32com.client

    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(docx_path.resolve()))
        try:
            doc.Fields.Update()
            for toc in doc.TablesOfContents:
                toc.Update()
        except Exception:
            pass

        doc.Save()
        doc.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
        doc.Close(False)
    finally:
        word.Quit()

def set_paragraph_text_keep_first_run(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ''

def update_cover_date(cover_docx_path, display_date):
    doc = Document(cover_docx_path)
    changed = False
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text:
            continue

        if re.search(r'Last Updated:', text, re.IGNORECASE):
            set_paragraph_text_keep_first_run(paragraph, f'Last Updated: {display_date}')
            changed = True
            continue

        if re.search(r'Cập nhật lần cuối', text, re.IGNORECASE):
            set_paragraph_text_keep_first_run(paragraph, f'(Cập nhật lần cuối {display_date})')
            changed = True

    if changed:
        doc.save(cover_docx_path)

def prepare_cover_pdf(root_dir, display_date):
    cover_docx_path = root_dir / 'word' / 'bia' / 'so_tay_a5_bia.docx'
    cover_pdf_path = root_dir / 'pdf' / 'bia' / 'so_tay_a5_bia.pdf'
    cover_pdf_path.parent.mkdir(parents=True, exist_ok=True)

    if cover_docx_path.exists():
        update_cover_date(cover_docx_path, display_date)
        update_docx_fields_and_export_pdf(cover_docx_path, cover_pdf_path)
        print(f"Đã cập nhật ngày bìa và xuất PDF bìa: {cover_pdf_path}")
        return cover_pdf_path

    return cover_pdf_path

def get_display_date():
    return datetime.now().strftime("%d/%m/%Y")

def merge_pdfs(input_paths, output_path):
    from pypdf import PdfWriter

    logging.getLogger('pypdf').setLevel(logging.ERROR)
    writer = PdfWriter()
    for input_path in input_paths:
        writer.append(str(input_path))

    with open(output_path, 'wb') as f:
        writer.write(f)

def ensure_output_dirs(root_dir):
    word_dir = root_dir / 'word' / 'sotay'
    pdf_dir = root_dir / 'pdf' / 'sotay'
    word_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir.mkdir(parents=True, exist_ok=True)
    return word_dir, pdf_dir

def add_section_with_break(doc, add_func, *args):
    if len(doc.paragraphs) > 0 or len(doc.tables) > 0:
        doc.add_page_break()
    add_func(doc, *args)

# --- HÀM CHÍNH ---
def tao_tai_lieu_mau():
    choices = ask_export_options()
    root_dir = PROJECT_ROOT
    display_date = get_display_date()
    word_dir, pdf_dir = ensure_output_dirs(root_dir)
    doc = setup_document()

    # ================= BƯỚC CHUẨN BỊ DỮ LIỆU =================
    tuvung_dir = os.path.join(root_dir, 'tuvung')
    all_words, json_data_map = load_vocabulary_data(tuvung_dir)
    five_s_data = load_5s_data(tuvung_dir)

    add_toc_section(doc)
    add_section_with_break(doc, add_5s_section, five_s_data)

    if '1' in choices:
        add_section_with_break(doc, add_japanese_index, all_words)

    if '2' in choices:
        add_section_with_break(doc, add_vietnamese_index, all_words)

    if '3' in choices:
        add_section_with_break(doc, add_grouped_vocabulary, json_data_map)

    qr_path = root_dir / 'assets' / 'qr' / 'so_tay_a5_online_qr.png'
    add_section_with_break(doc, add_online_version_page, qr_path, display_date)

    docx_path = word_dir / 'so_tay_a5_no_bia.docx'
    content_pdf_path = pdf_dir / 'so_tay_a5_no_bia.pdf'
    final_pdf_path = pdf_dir / 'so_tay_a5.pdf'
    cover_pdf_path = prepare_cover_pdf(root_dir, display_date)

    doc.save(docx_path)
    print(f"Đã tạo file Word: {docx_path}")

    update_docx_fields_and_export_pdf(docx_path, content_pdf_path)
    print(f"Đã tạo PDF nội dung: {content_pdf_path}")

    if cover_pdf_path.exists():
        merge_pdfs([cover_pdf_path, content_pdf_path], final_pdf_path)
        print(f"Đã ghép bìa và tạo PDF hoàn chỉnh: {final_pdf_path}")
    else:
        print(f"Không tìm thấy file bìa: {cover_pdf_path}")

if __name__ == "__main__":
    tao_tai_lieu_mau()

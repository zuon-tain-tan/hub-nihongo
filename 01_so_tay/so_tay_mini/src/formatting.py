from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from config import (
    BOTTOM_MARGIN_CM,
    DOCUMENT_FONT_SIZE_PT,
    FOOTER_DISTANCE_CM,
    GUTTER_CM,
    JAPANESE_FONT,
    LEFT_MARGIN_CM,
    PAGE_HEIGHT_CM,
    PAGE_WIDTH_CM,
    RIGHT_MARGIN_CM,
    TOP_MARGIN_CM,
    VIETNAMESE_FONT,
)
from docx_utils import add_page_number, apply_paragraph_format, apply_run_font


def setup_document(footer_text="Sổ tay từ vựng N5 - N4", booklet_layout=True):
    doc = Document()
    setup_page(doc, booklet_layout=booklet_layout)
    setup_styles(doc)
    setup_footer(doc, footer_text, booklet_layout=booklet_layout)
    return doc


def setup_page(doc, booklet_layout=True):
    section = doc.sections[0]
    section.page_width = Cm(PAGE_WIDTH_CM)
    section.page_height = Cm(PAGE_HEIGHT_CM)
    section.top_margin = Cm(TOP_MARGIN_CM)
    section.bottom_margin = Cm(BOTTOM_MARGIN_CM)
    section.left_margin = Cm(LEFT_MARGIN_CM)
    section.right_margin = Cm(RIGHT_MARGIN_CM if booklet_layout else LEFT_MARGIN_CM)
    section.gutter = Cm(GUTTER_CM if booklet_layout else 0)
    section.footer_distance = Cm(FOOTER_DISTANCE_CM)


def setup_styles(doc):
    configure_style(doc.styles["Normal"])

    table_style = doc.styles.add_style("Bang_Style", 1)
    configure_style(table_style)

    index_style = doc.styles.add_style("Index_Style", 1)
    configure_style(index_style)
    index_style.paragraph_format.tab_stops.add_tab_stop(
        Cm(11.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )

    for level in range(1, 4):
        heading_style = doc.styles[f"Heading {level}"]
        configure_style(heading_style, color=RGBColor(0, 0, 0))
        heading_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        heading_style.paragraph_format.space_before = Pt(0)
        heading_style.paragraph_format.space_after = Pt(0)
        heading_style.paragraph_format.keep_with_next = True
        heading_style.paragraph_format.keep_together = True
        if level == 1:
            heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for toc_style_name in ["TOC 1", "TOC 2", "TOC 3"]:
        try:
            configure_style(doc.styles[toc_style_name])
        except KeyError:
            continue


def configure_style(style, color=None):
    font = style.font
    font.name = VIETNAMESE_FONT
    font.size = Pt(DOCUMENT_FONT_SIZE_PT)
    if color is not None:
        font.color.rgb = color

    r_fonts = font.element.rPr.rFonts
    if r_fonts is not None:
        for theme_attr in ["w:asciiTheme", "w:hAnsiTheme", "w:cstheme"]:
            if qn(theme_attr) in r_fonts.attrib:
                del r_fonts.attrib[qn(theme_attr)]
        r_fonts.set(qn("w:eastAsia"), JAPANESE_FONT)

    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)


def setup_footer(doc, footer_text, booklet_layout=True):
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = booklet_layout

    if not booklet_layout:
        footer_paragraph = section.footer.paragraphs[0]
        apply_paragraph_format(footer_paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
        apply_run_font(footer_paragraph.add_run(f"{footer_text} | "))
        page_run = footer_paragraph.add_run()
        add_page_number(page_run)
        apply_run_font(page_run)
        return

    even_footer_paragraph = section.even_page_footer.paragraphs[0]
    apply_paragraph_format(even_footer_paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    even_page_run = even_footer_paragraph.add_run()
    add_page_number(even_page_run)
    apply_run_font(even_page_run)
    apply_run_font(even_footer_paragraph.add_run(f" | {footer_text}"))

    odd_footer_paragraph = section.footer.paragraphs[0]
    apply_paragraph_format(odd_footer_paragraph, WD_ALIGN_PARAGRAPH.RIGHT)
    apply_run_font(odd_footer_paragraph.add_run(f"{footer_text} | "))
    odd_page_run = odd_footer_paragraph.add_run()
    add_page_number(odd_page_run)
    apply_run_font(odd_page_run)

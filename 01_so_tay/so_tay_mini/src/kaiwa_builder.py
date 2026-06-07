from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from config import DATA_DIR, OUTPUT_DIR
from data_loader import load_kaiwa_before_departure
from docx_utils import apply_paragraph_format, apply_run_font, keep_with_next
from export import update_docx_fields_and_export_pdf
from formatting import setup_document


KAIWA_DOCX_NAME = "kaiwa_truoc_xuat_canh.docx"
KAIWA_PDF_NAME = "kaiwa_truoc_xuat_canh.pdf"


def setup_kaiwa_document():
    doc = setup_document(footer_text="Kaiwa trước xuất cảnh", booklet_layout=False)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.gutter = Cm(0)
    return doc


def add_centered_run_paragraph(doc, text, bold=False):
    paragraph = doc.add_paragraph()
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(paragraph.add_run(text), bold=bold)
    return paragraph


def add_body_paragraph(doc, text, bold=False, keep_next=False):
    paragraph = doc.add_paragraph(style="Bang_Style")
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    apply_run_font(paragraph.add_run(text), bold=bold)
    if keep_next:
        keep_with_next(paragraph)
    return paragraph


def add_kaiwa_qa(doc, qa, question_number):
    add_body_paragraph(doc, qa.get("question_jp", ""), bold=True, keep_next=True)
    add_body_paragraph(doc, qa.get("question_vi", ""), keep_next=True)
    add_body_paragraph(doc, qa.get("answer_jp", ""), keep_next=True)
    add_body_paragraph(doc, qa.get("answer_vi", ""))


def add_kaiwa_content(doc, data):
    title = add_centered_run_paragraph(doc, data.get("title", "KAIWA TRƯỚC XUẤT CẢNH"), bold=True)
    keep_with_next(title)
    subtitle = data.get("subtitle")
    if subtitle:
        keep_with_next(add_centered_run_paragraph(doc, subtitle))

    description = data.get("description")
    if description:
        add_centered_run_paragraph(doc, description)

    question_number = 1
    for section_index, section in enumerate(data.get("sections", []), start=1):
        heading = doc.add_heading(f"PHẦN {section_index}. {section.get('title_vi', '')}", level=1)
        keep_with_next(heading)
        if section_index > 1:
            heading.paragraph_format.page_break_before = True

        for item_index, item in enumerate(section.get("items", []), start=1):
            item_heading = doc.add_heading(f"{question_number}. {item.get('title_vi', '')}", level=2)
            keep_with_next(item_heading)
            for qa in item.get("qa", []):
                add_kaiwa_qa(doc, qa, question_number)
                question_number += 1
            if item_index < len(section.get("items", [])):
                add_body_paragraph(doc, "")

    advice = data.get("advice", [])
    if advice:
        heading = doc.add_heading("LỜI KHUYÊN KHI LUYỆN KAIWA", level=1)
        heading.paragraph_format.page_break_before = True
        keep_with_next(heading)
        for index, text in enumerate(advice, start=1):
            add_body_paragraph(doc, f"{index}. {text}")


def create_kaiwa_document(data):
    doc = setup_kaiwa_document()
    add_kaiwa_content(doc, data)
    return doc


def ensure_kaiwa_output_dir():
    docx_dir = OUTPUT_DIR / "docx" / "kaiwa"
    pdf_dir = OUTPUT_DIR / "pdf" / "kaiwa"
    for output_dir in [docx_dir, pdf_dir]:
        output_dir.mkdir(parents=True, exist_ok=True)
    return docx_dir, pdf_dir


def build_kaiwa_before_departure_document():
    data = load_kaiwa_before_departure(DATA_DIR)
    doc = create_kaiwa_document(data)
    docx_dir, pdf_dir = ensure_kaiwa_output_dir()
    docx_path = docx_dir / KAIWA_DOCX_NAME
    pdf_path = pdf_dir / KAIWA_PDF_NAME
    doc.save(docx_path)
    print(f"Đã tạo file Word Kaiwa: {docx_path}")
    update_docx_fields_and_export_pdf(docx_path, pdf_path)
    print(f"Đã tạo PDF Kaiwa: {pdf_path}")
    return docx_path, pdf_path

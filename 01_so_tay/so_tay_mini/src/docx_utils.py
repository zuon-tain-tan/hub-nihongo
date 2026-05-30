from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from config import (
    DOCUMENT_FONT_SIZE_PT,
    JAPANESE_FONT,
    JAPANESE_FONT_SIZE_PT,
    LINE_SPACING_TWIPS,
    TABLE_CELL_VERTICAL_MARGIN_TWIPS,
    TABLE_LINE_SPACING_TWIPS,
    VIETNAMESE_FONT,
)


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(qn(name), value)


def add_page_number(run):
    field_begin = create_element("w:fldChar")
    create_attribute(field_begin, "w:fldCharType", "begin")

    instruction = create_element("w:instrText")
    create_attribute(instruction, "xml:space", "preserve")
    instruction.text = "PAGE"

    field_separate = create_element("w:fldChar")
    create_attribute(field_separate, "w:fldCharType", "separate")

    field_end = create_element("w:fldChar")
    create_attribute(field_end, "w:fldCharType", "end")

    run._r.extend([field_begin, instruction, field_separate, field_end])


def add_simple_field(paragraph, instruction):
    run = paragraph.add_run()
    field_begin = OxmlElement("w:fldChar")
    field_begin.set(qn("w:fldCharType"), "begin")

    instruction_text = OxmlElement("w:instrText")
    instruction_text.set(qn("xml:space"), "preserve")
    instruction_text.text = instruction

    field_separate = OxmlElement("w:fldChar")
    field_separate.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = "Bấm Ctrl+A rồi F9 để cập nhật mục lục nếu cần."

    field_end = OxmlElement("w:fldChar")
    field_end.set(qn("w:fldCharType"), "end")

    run._r.extend([field_begin, instruction_text, field_separate, placeholder, field_end])


def contains_japanese(text):
    return any(
        "\u3040" <= char <= "\u30ff" or "\u3400" <= char <= "\u9fff"
        for char in text
    )


def get_run_font_size(text, size=None):
    if size is not None:
        return size
    if contains_japanese(text):
        return JAPANESE_FONT_SIZE_PT
    return DOCUMENT_FONT_SIZE_PT


def apply_run_font(run, bold=False, size=None):
    run.font.bold = bold
    run.font.size = Pt(get_run_font_size(run.text, size))
    run.font.name = VIETNAMESE_FONT
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), JAPANESE_FONT)


def apply_paragraph_format(paragraph, alignment=None):
    if alignment is not None:
        paragraph.alignment = alignment

    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_paragraph_spacing(paragraph, before=0, after=0, line=LINE_SPACING_TWIPS)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True
    return paragraph


def apply_table_paragraph_format(paragraph, alignment=None):
    if alignment is not None:
        paragraph.alignment = alignment

    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    set_paragraph_spacing(paragraph, before=0, after=0, line=TABLE_LINE_SPACING_TWIPS)


def set_table_cell_margins(table, top=None, bottom=None):
    table_properties = table._tbl.tblPr
    margins = table_properties.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        table_properties.append(margins)

    for side, value in {"top": top, "bottom": bottom}.items():
        if value is None:
            continue
        margin = margins.find(qn(f"w:{side}"))
        if margin is None:
            margin = OxmlElement(f"w:{side}")
            margins.append(margin)
        margin.set(qn("w:w"), str(value))
        margin.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    cell.width = width
    cell_properties = cell._tc.get_or_add_tcPr()
    cell_width = cell_properties.first_child_found_in("w:tcW")
    if cell_width is None:
        cell_width = OxmlElement("w:tcW")
        cell_properties.append(cell_width)
    cell_width.set(qn("w:w"), str(int(width.cm * 567)))
    cell_width.set(qn("w:type"), "dxa")


def set_cell_shading(cell, fill):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = cell_properties.first_child_found_in("w:shd")
    if shading is None:
        shading = OxmlElement("w:shd")
        cell_properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_paragraph_spacing(paragraph, before=0, after=0, line=LINE_SPACING_TWIPS, line_rule="auto"):
    paragraph_properties = paragraph._p.get_or_add_pPr()
    spacing = paragraph_properties.first_child_found_in("w:spacing")
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        paragraph_properties.append(spacing)

    spacing.set(qn("w:before"), str(before))
    spacing.set(qn("w:after"), str(after))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), line_rule)


def add_table(doc, headers, rows, column_widths=None):
    header_count = 1 if headers else 0
    column_count = len(headers) if headers else len(rows[0]) if rows else 1
    table = doc.add_table(rows=header_count, cols=column_count)
    table.style = "Table Grid"
    table.autofit = False
    set_table_cell_margins(
        table,
        top=TABLE_CELL_VERTICAL_MARGIN_TWIPS,
        bottom=TABLE_CELL_VERTICAL_MARGIN_TWIPS,
    )

    if headers:
        header_cells = table.rows[0].cells
        for index, header in enumerate(headers):
            header_cells[index].text = header
            if column_widths:
                set_cell_width(header_cells[index], column_widths[index])

        for cell in header_cells:
            set_cell_shading(cell, "D9D9D9")
            format_table_cell(cell)

    for row in rows:
        row_cells = table.add_row().cells
        for index, value in enumerate(row):
            row_cells[index].text = value
            if column_widths:
                set_cell_width(row_cells[index], column_widths[index])

        for cell in row_cells:
            format_table_cell(cell)


def format_table_cell(cell, bold=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for paragraph in cell.paragraphs:
        paragraph.style = "Bang_Style"
        apply_table_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
        for run in paragraph.runs:
            apply_run_font(run, bold=bold)


def set_paragraph_text_keep_first_run(paragraph, text):
    if not paragraph.runs:
        paragraph.add_run(text)
        return

    paragraph.runs[0].text = text
    for run in paragraph.runs[1:]:
        run.text = ""

import json
import re
import unicodedata
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from config import ASSETS_DIR, DATA_DIR, OUTPUT_DIR, PROJECT_ROOT
from data_loader import load_vocabulary_data
from docx_utils import apply_paragraph_format, apply_run_font, keep_with_next
from export import update_docx_fields_and_export_pdf
from formatting import setup_document
from image_vocabulary_scraper import (
    download_langoal_vocabulary_lesson,
    format_lesson_dir_name,
)


IMAGE_VOCABULARY_DOCX_NAME_TEMPLATE = "hinh_anh_tu_vung_lesson_{lesson:02d}.docx"
IMAGE_VOCABULARY_PDF_NAME_TEMPLATE = "hinh_anh_tu_vung_lesson_{lesson:02d}.pdf"
IMAGE_VOCABULARY_FULL_DOCX_NAME_TEMPLATE = "hinh_anh_tu_vung_{selection}.docx"
IMAGE_VOCABULARY_FULL_PDF_NAME_TEMPLATE = "hinh_anh_tu_vung_{selection}.pdf"
MIN_LESSON = 1
MAX_LESSON = 50
JAPANESE_QUOTE_PATTERN = re.compile(r"「([^」]+)」")
MANUAL_READING_OVERRIDES = {
    "鍵": "かぎ",
    "椅子": "いす",
    "煙草": "たばこ",
    "一人で": "ひとりで",
    "牛肉": "ぎゅうにく",
    "豚肉": "ぶたにく",
    "CD": "シーディー",
    "貸します": "かします",
    "借ります": "かります",
    "静かな": "しずかな",
    "有名な": "ゆうめいな",
    "親切な": "しんせつな",
    "元気な": "げんきな",
    "暇な": "ひまな",
    "便利な": "べんりな",
    "暑い": "あつい",
    "熱い": "あつい",
    "車": "くるま",
    "勉強": "べんきょう",
    "好きな": "すきな",
    "嫌いな": "きらいな",
    "上手な": "じょうずな",
    "下手な": "へたな",
    "色々な": "いろいろな",
    "兄弟": "きょうだい",
    "兄": "あに",
    "お兄さん": "おにいさん",
    "姉": "あね",
    "お姉さん": "おねえさん",
    "弟": "おとうと",
    "弟さん": "おとうとさん",
    "妹": "いもうと",
    "妹さん": "いもうとさん",
    "早い": "はやい",
    "速い": "はやい",
    "すき焼き": "すきやき",
    "寿司": "すし",
    "天ぷら": "てんぷら",
    "大変な": "たいへんな",
    "お腹がすきました。": "おなかがすきました",
    "喉が渇きました。": "のどがかわきました",
    "駐めます": "とめます",
    "口": "くち",
    "お腹": "おなか",
    "大切な": "たいせつな",
    "大丈夫な": "だいじょうぶな",
    "風呂": "ふろ",
    "換えます": "かえます",
    "不便な": "ふべんな",
    "気をつけます": "きをつけます",
    "連れて行きます": "つれていきます",
    "連れて来ます": "つれてきます",
    "淹れます": "いれます",
    "弁当": "べんとう",
    "参加します": "さんかします",
    "申し込みます": "もうしこみます",
    "都合がいい": "つごうがいい",
    "都合が悪い": "つごうがわるい",
    "柔道": "じゅうどう",
    "湯": "ゆ",
    "皿": "さら",
    "茶碗": "ちゃわん",
    "出席します": "しゅっせきします",
    "治ります": "なおります",
    "直ります": "なおります",
    "怪我": "けが",
    "咳": "せき",
    "諦めます": "あきらめます",
    "柔らかい": "やわらかい",
    "叱ります": "しかります",
    "嘘": "うそ",
    "判子": "はんこ",
    "遅刻します": "ちこくします",
    "早退します": "そうたいします",
    "喧嘩します": "けんかします",
    "嬉しい": "うれしい",
    "見合い": "みあい",
    "号室": "ごうしつ",
    "汗": "あせ",
    "石鹸": "せっけん",
    "申し込み": "もうしこみ",
    "見舞い": "みまい",
    "餌": "えさ",
    "指輪": "ゆびわ",
    "預かります": "あずかります",
    "助かります": "たすかります",
    "石": "いし",
    "取れます": "とれます",
    "適当な": "てきとうな",
    "年齢": "ねんれい",
    "収入": "しゅうにゅう",
    "丁寧": "ていねい",
    "丁寧な": "ていねいな",
    "細かい": "こまかい",
    "倍": "ばい",
    "箪笥": "たんす",
    "贈り物": "おくりもの",
    "間違い電話": "まちがいでんわ",
    "札": "さつ",
    "注射": "ちゅうしゃ",
    "食欲": "しょくよく",
    "長生きします": "ながいきします",
    "発表": "はっぴょう",
    "下ろします": "おろします",
    "久しぶり": "ひさしぶり",
    "寄ります": "よります",
    "貿易": "ぼうえき",
    "放送します": "ほうそうします",
    "象": "ぞう",
}


def setup_image_vocabulary_document():
    doc = setup_document(footer_text="Hình ảnh từ vựng", booklet_layout=False)
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.gutter = Cm(0)
    return doc


def add_centered_paragraph(doc, text, bold=False, size=None):
    paragraph = doc.add_paragraph()
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(paragraph.add_run(text), bold=bold, size=size)
    return paragraph


def add_source_paragraph(doc, source_url):
    paragraph = doc.add_paragraph()
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(paragraph.add_run(f"Nguồn: {source_url}"), size=10)
    return paragraph


def set_cell_margins(cell):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(6)


def add_image_vocabulary_card(cell, item, project_root):
    set_cell_margins(cell)
    image_path = project_root / item["image_file"]
    image_paragraph = cell.paragraphs[0]
    apply_paragraph_format(image_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    if image_path.exists():
        image_paragraph.add_run().add_picture(str(image_path), width=Cm(4.8))
    else:
        apply_run_font(image_paragraph.add_run("[Thiếu ảnh]"), bold=True)

    title_paragraph = cell.add_paragraph()
    apply_paragraph_format(title_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(title_paragraph.add_run(item.get("vocabulary", "")), bold=True)

    reading = item.get("reading", "")
    if reading and reading != item.get("vocabulary", ""):
        reading_paragraph = cell.add_paragraph()
        apply_paragraph_format(reading_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        apply_run_font(reading_paragraph.add_run(reading), size=11)


def add_image_vocabulary_grid(doc, items, project_root, columns=3):
    table = doc.add_table(rows=0, cols=columns)
    table.autofit = True

    for start in range(0, len(items), columns):
        row = table.add_row()
        row_items = items[start : start + columns]
        for column_index, item in enumerate(row_items):
            add_image_vocabulary_card(row.cells[column_index], item, project_root)

        for column_index in range(len(row_items), columns):
            row.cells[column_index].text = ""

    return table


def add_image_vocabulary_lesson(doc, metadata, project_root=PROJECT_ROOT):
    title = add_centered_paragraph(
        doc,
        f"HÌNH ẢNH TỪ VỰNG - BÀI {metadata['lesson']}",
        bold=True,
        size=16,
    )
    keep_with_next(title)
    add_source_paragraph(doc, metadata.get("source_url", ""))
    doc.add_paragraph()
    add_image_vocabulary_grid(doc, metadata.get("items", []), project_root)


def create_image_vocabulary_document(metadata, project_root=PROJECT_ROOT):
    doc = setup_image_vocabulary_document()
    add_image_vocabulary_lesson(doc, metadata, project_root)
    return doc


def create_image_vocabulary_full_document(metadata_list, project_root=PROJECT_ROOT):
    doc = setup_image_vocabulary_document()
    for index, metadata in enumerate(metadata_list):
        if index > 0:
            doc.add_page_break()
        add_image_vocabulary_lesson(doc, metadata, project_root)
    return doc


def ensure_image_vocabulary_output_dirs():
    docx_dir = OUTPUT_DIR / "docx" / "image_vocabulary"
    pdf_dir = OUTPUT_DIR / "pdf" / "image_vocabulary"
    for output_dir in [docx_dir, pdf_dir]:
        output_dir.mkdir(parents=True, exist_ok=True)
    return docx_dir, pdf_dir


def load_image_vocabulary_metadata(lesson, data_dir=DATA_DIR):
    path = data_dir / "image_vocabulary" / f"{format_lesson_dir_name(lesson)}.json"
    return json.loads(path.read_text(encoding="utf-8"))


def get_image_vocabulary_metadata_path(lesson, data_dir=DATA_DIR):
    return data_dir / "image_vocabulary" / f"{format_lesson_dir_name(lesson)}.json"


def normalize_japanese_lookup_text(text):
    text = unicodedata.normalize("NFKC", text or "")
    text = text.strip()
    text = re.sub(r"[\[［][^\]］]*[\]］]", "", text)
    text = text.strip("「」『』[]［］")
    text = text.replace("～", "").replace("~", "")
    text = re.sub(r"\s+", "", text)
    return text


def get_lookup_candidates(text):
    normalized = normalize_japanese_lookup_text(text)
    if not normalized:
        return []

    candidates = [normalized]
    without_optional_o = normalized.replace("(お)", "").replace("（お）", "")
    candidates.append(without_optional_o)
    candidates.append(normalized.replace("(お)", "お").replace("（お）", "お"))

    without_parentheses = re.sub(r"[（(][^()（）]*[）)]", "", normalized)
    candidates.append(without_parentheses)
    candidates.append(normalized.lstrip("ー-－"))

    cleaned = []
    seen = set()
    for candidate in candidates:
        candidate = normalize_japanese_lookup_text(candidate)
        if candidate and candidate not in seen:
            seen.add(candidate)
            cleaned.append(candidate)
    return cleaned


def contains_kana_or_kanji(text):
    return any(
        "\u3040" <= char <= "\u30ff" or "\u3400" <= char <= "\u9fff"
        for char in text or ""
    )


def split_langoal_vocabulary_terms(text):
    terms = JAPANESE_QUOTE_PATTERN.findall(text or "")
    if not terms and text:
        terms = re.split(r"[/／,，、]+", text)
    normalized_terms = []
    for term in terms:
        normalized_terms.extend(get_lookup_candidates(term))
    return [term for index, term in enumerate(normalized_terms) if term and term not in normalized_terms[:index]]


def split_local_vocabulary_text(text):
    text = unicodedata.normalize("NFKC", text or "").strip()
    match = re.match(r"^(.*?)\s*[（(]([^()（）]+)[）)]\s*$", text)
    if not match:
        normalized = normalize_japanese_lookup_text(text)
        return [(normalized, normalized)] if normalized else []

    reading, written = match.groups()
    reading = normalize_japanese_lookup_text(reading)
    written = normalize_japanese_lookup_text(written)
    if not reading or not written:
        return []

    return [(written, reading), (reading, reading)]


def add_reading_index_entry(reading_index, key, reading):
    reading = normalize_japanese_lookup_text(reading)
    for candidate in get_lookup_candidates(key):
        if not candidate or not reading:
            continue
        reading_index.setdefault(candidate, reading)


def build_vocabulary_reading_index(data_dir=DATA_DIR):
    all_words, _json_data_map = load_vocabulary_data(data_dir)
    reading_index = {}
    for item in all_words:
        for key, reading in split_local_vocabulary_text(item.get("tu_vung", "")):
            add_reading_index_entry(reading_index, key, reading)
    for key, reading in MANUAL_READING_OVERRIDES.items():
        add_reading_index_entry(reading_index, key, reading)
    return reading_index


def get_term_reading(term, reading_index):
    for candidate in get_lookup_candidates(term):
        if candidate in reading_index:
            return reading_index[candidate]
        if contains_kana_or_kanji(candidate) and not any("\u3400" <= char <= "\u9fff" for char in candidate):
            return candidate
    return ""


def get_vocabulary_reading(vocabulary, reading_index):
    readings = []
    seen = set()
    for term in split_langoal_vocabulary_terms(vocabulary):
        reading = get_term_reading(term, reading_index)
        if reading and reading not in seen:
            seen.add(reading)
            readings.append(reading)
    return " / ".join(readings)


def enrich_image_vocabulary_metadata(metadata, reading_index=None):
    reading_index = reading_index if reading_index is not None else build_vocabulary_reading_index()
    changed = False
    for item in metadata.get("items", []):
        reading = get_vocabulary_reading(item.get("vocabulary", ""), reading_index)
        if reading and item.get("reading") != reading:
            item["reading"] = reading
            changed = True
    return changed


def save_image_vocabulary_metadata(metadata, lesson, data_dir=DATA_DIR):
    path = get_image_vocabulary_metadata_path(lesson, data_dir)
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(metadata, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )
    return path


def parse_lesson_number(raw):
    raw = unicodedata.normalize("NFKC", raw or "").strip()
    if not raw.isdigit():
        return None
    lesson = int(raw)
    if lesson < MIN_LESSON or lesson > MAX_LESSON:
        return None
    return lesson


def parse_lesson_selection(raw):
    raw = unicodedata.normalize("NFKC", raw or "").strip()
    if not raw:
        return None

    lessons = []
    seen = set()
    raw = re.sub(r"\s*-\s*", "-", raw)
    tokens = [token for token in re.split(r"[\s,;、]+", raw) if token]
    for token in tokens:
        if "-" in token:
            parts = token.split("-", 1)
            if len(parts) != 2 or not parts[0].isdigit() or not parts[1].isdigit():
                return None
            start, end = int(parts[0]), int(parts[1])
            if start > end:
                return None
            token_lessons = range(start, end + 1)
        else:
            if not token.isdigit():
                return None
            token_lessons = [int(token)]

        for lesson in token_lessons:
            if lesson < MIN_LESSON or lesson > MAX_LESSON:
                return None
            if lesson not in seen:
                seen.add(lesson)
                lessons.append(lesson)

    return lessons or None


def format_lesson_selection_slug(lessons):
    lessons = list(lessons)
    if len(lessons) == 1:
        return f"lesson_{lessons[0]:02d}"
    if lessons == list(range(lessons[0], lessons[-1] + 1)):
        return f"lessons_{lessons[0]:02d}_{lessons[-1]:02d}"
    return "lessons_" + "_".join(f"{lesson:02d}" for lesson in lessons)


def ask_image_vocabulary_lessons():
    while True:
        raw = input("Nhập bài muốn tải, ví dụ 1 hoặc 1-50 hoặc 1,3,5: ")
        lessons = parse_lesson_selection(raw)
        if lessons is not None:
            return lessons
        print("Lựa chọn không hợp lệ. Vui lòng nhập trong khoảng 1-50, ví dụ: 1 hoặc 1-50")


def get_or_download_image_vocabulary_metadata(lesson):
    metadata_path = get_image_vocabulary_metadata_path(lesson)
    reading_index = build_vocabulary_reading_index(DATA_DIR)
    if metadata_path.exists():
        print(f"Đã có metadata lesson-{lesson}, bỏ qua tải lại: {metadata_path}")
        metadata = load_image_vocabulary_metadata(lesson)
        if enrich_image_vocabulary_metadata(metadata, reading_index):
            save_image_vocabulary_metadata(metadata, lesson)
            print(f"Đã bổ sung cách đọc vào metadata: {metadata_path}")
        return metadata, metadata_path, False

    print(f"Đang mở Langoal lesson-{lesson} và tải ảnh từ vựng...")
    metadata, metadata_path = download_langoal_vocabulary_lesson(
        lesson,
        ASSETS_DIR,
        DATA_DIR,
        PROJECT_ROOT,
    )
    enrich_image_vocabulary_metadata(metadata, reading_index)
    save_image_vocabulary_metadata(metadata, lesson)
    print(f"Đã lưu metadata: {metadata_path}")
    return metadata, metadata_path, True


def get_image_vocabulary_output_paths(lessons, docx_dir, pdf_dir):
    lessons = list(lessons)
    if len(lessons) == 1:
        lesson = lessons[0]
        return (
            docx_dir / IMAGE_VOCABULARY_DOCX_NAME_TEMPLATE.format(lesson=lesson),
            pdf_dir / IMAGE_VOCABULARY_PDF_NAME_TEMPLATE.format(lesson=lesson),
        )

    selection = format_lesson_selection_slug(lessons)
    return (
        docx_dir / IMAGE_VOCABULARY_FULL_DOCX_NAME_TEMPLATE.format(selection=selection),
        pdf_dir / IMAGE_VOCABULARY_FULL_PDF_NAME_TEMPLATE.format(selection=selection),
    )


def build_image_vocabulary_document(lessons=None):
    if lessons is None:
        lessons = ask_image_vocabulary_lessons()
    elif isinstance(lessons, int):
        lessons = [lessons]
    else:
        lessons = list(lessons)

    metadata_list = []
    for lesson in lessons:
        metadata, _metadata_path, _downloaded = get_or_download_image_vocabulary_metadata(lesson)
        metadata_list.append(metadata)

    if len(metadata_list) == 1:
        doc = create_image_vocabulary_document(metadata_list[0])
    else:
        doc = create_image_vocabulary_full_document(metadata_list)

    docx_dir, pdf_dir = ensure_image_vocabulary_output_dirs()
    docx_path, pdf_path = get_image_vocabulary_output_paths(lessons, docx_dir, pdf_dir)
    doc.save(docx_path)
    print(f"Đã tạo file Word hình ảnh từ vựng: {docx_path}")
    update_docx_fields_and_export_pdf(docx_path, pdf_path)
    print(f"Đã tạo PDF hình ảnh từ vựng: {pdf_path}")
    return docx_path, pdf_path

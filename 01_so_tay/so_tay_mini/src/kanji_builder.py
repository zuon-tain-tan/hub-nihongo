import random
import re

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from config import DATA_DIR, OUTPUT_DIR
from data_loader import load_vocabulary_data
from docx_utils import add_table, apply_paragraph_format, apply_run_font, keep_with_next
from export import update_docx_fields_and_export_pdf
from formatting import setup_document


KANJI_TABLE_WIDTHS = {
    "1": [Cm(3.4), Cm(4.1), Cm(4.0)],
    "2": [Cm(4.1), Cm(3.4), Cm(4.0)],
    "3": [Cm(4.4), Cm(3.4), Cm(3.7)],
}
MIXED_ANSWER_WIDTHS = [Cm(1.0), Cm(1.2), Cm(2.7), Cm(3.0), Cm(3.6)]

KANJI_MODES = {
    "1": {
        "name": "kieu1",
        "title": "Bản có Kanji, trống ô mềm, trống nghĩa",
        "headers": ["Kanji", "Chữ mềm", "Nghĩa"],
        "row": lambda item, answer=False: [
            item["kanji"],
            item["kana"] if answer else "",
            item["meaning"] if answer else "",
        ],
        "widths": KANJI_TABLE_WIDTHS["1"],
    },
    "2": {
        "name": "kieu2",
        "title": "Bản có chữ mềm, trống ô Kanji, trống nghĩa",
        "headers": ["Chữ mềm", "Kanji", "Nghĩa"],
        "row": lambda item, answer=False: [
            item["kana"],
            item["kanji"] if answer else "",
            item["meaning"] if answer else "",
        ],
        "widths": KANJI_TABLE_WIDTHS["2"],
    },
    "3": {
        "name": "kieu3",
        "title": "Bản có nghĩa, trống ô Kanji, trống mềm",
        "headers": ["Nghĩa", "Kanji", "Chữ mềm"],
        "row": lambda item, answer=False: [
            item["meaning"],
            item["kanji"] if answer else "",
            item["kana"] if answer else "",
        ],
        "widths": KANJI_TABLE_WIDTHS["3"],
    },
}


def ask_choice(prompt, valid_options, default=None):
    while True:
        raw = input(prompt).strip().lower()
        if not raw and default is not None:
            return default
        if raw in valid_options:
            return raw
        print(f"Lựa chọn không hợp lệ. Vui lòng nhập: {', '.join(valid_options)}.")


def ask_kanji_levels():
    print("Chọn cấp độ Kanji:")
    print("1. Cấp độ N5")
    print("2. Cấp độ N4")
    print("3. Cả 2 cấp N5, N4")
    choice = ask_choice("Nhập lựa chọn: ", {"1", "2", "3"}, default="3")
    if choice == "1":
        return {"N5"}
    if choice == "2":
        return {"N4"}
    return {"N5", "N4"}


def ask_kanji_mode():
    print("Chọn loại đề:")
    for key, mode in KANJI_MODES.items():
        print(f"{key}. {mode['title']}")
    choice = ask_choice("Nhập lựa chọn: ", set(KANJI_MODES), default="1")
    return KANJI_MODES[choice]


def ask_kanji_build_type():
    print("Chọn cách tạo Kanji:")
    print("1. Tạo Kanji theo bài, kèm đáp án ở cuối")
    print("2. Tạo Kanji lộn xộn, kèm đáp án ở trang cuối")
    print("3. Tương tự 1, không đáp án")
    print("4. Tương tự 2, không đáp án")
    return ask_choice("Nhập lựa chọn: ", {"1", "2", "3", "4"}, default="1")


def parse_lesson_selection(raw, available_lessons):
    raw = re.sub(r"\s*-\s*", "-", raw.strip().lower())
    if raw in {"", "all", "a", "tat ca", "tất cả"}:
        return set(available_lessons)

    selected = set()
    for part in [part for part in re.split(r"[\s,;]+", raw) if part]:
        range_match = re.fullmatch(r"(\d+)\s*-\s*(\d+)", part)
        if range_match:
            start, end = [int(value) for value in range_match.groups()]
            if start > end:
                start, end = end, start
            selected.update(range(start, end + 1))
            continue

        if part.isdigit():
            selected.add(int(part))
            continue

        raise ValueError

    return selected & set(available_lessons)


def ask_lessons(available_lessons):
    lesson_text = ", ".join(str(value) for value in available_lessons)
    print(f"Các bài đang có dữ liệu Kanji: {lesson_text}")
    print("Nhập all để tạo toàn bộ, hoặc nhập dạng 26-30, hoặc 26,28,30.")
    while True:
        raw = input("Nhập số bài muốn tạo: ")
        try:
            selected = parse_lesson_selection(raw, available_lessons)
        except ValueError:
            selected = set()
        if selected:
            return selected
        print("Không tìm thấy bài hợp lệ. Ví dụ đúng: all hoặc 26-30 hoặc 26,28,30.")


def ask_mixed_count(total_count):
    while True:
        raw = input(f"Nhập số câu muốn trộn, hoặc all để lấy toàn bộ ({total_count} câu): ").strip().lower()
        if raw in {"", "all", "a", "tat ca", "tất cả"}:
            return total_count
        if raw.isdigit() and 0 < int(raw) <= total_count:
            return int(raw)
        print(f"Số câu không hợp lệ. Hãy nhập từ 1 đến {total_count}, hoặc all.")


def contains_kanji(text):
    return any("\u3400" <= char <= "\u9fff" for char in text or "")


def extract_kana_and_kanji(text):
    matches = re.findall(r"(.+?)\s*[\(（]([^()（）]+)[\)）]", text or "")
    for kana, kanji in matches:
        if contains_kanji(kanji):
            return kana.strip(), kanji.strip()
    return "", ""


def get_kanji_items(levels):
    _, json_data_map = load_vocabulary_data(DATA_DIR)
    items = []
    seen = set()

    for file_data in json_data_map:
        if file_data["level"] not in levels:
            continue

        lesson_number = file_data["prefix_num"]
        for item in file_data["data"]:
            kana, kanji = extract_kana_and_kanji(item.get("tu_vung", ""))
            if not kana or not kanji:
                continue

            kanji_item = {
                "level": file_data["level"],
                "lesson": lesson_number,
                "kana": kana,
                "kanji": kanji,
                "meaning": item.get("y_nghia", ""),
            }
            key = (
                kanji_item["level"],
                kanji_item["lesson"],
                kanji_item["kana"],
                kanji_item["kanji"],
                kanji_item["meaning"],
            )
            if key in seen:
                continue
            seen.add(key)
            items.append(kanji_item)

    return items


def add_title(doc, text):
    paragraph = doc.add_heading(text, level=1)
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    keep_with_next(paragraph)
    for run in paragraph.runs:
        apply_run_font(run, bold=True)


def add_subtitle(doc, text):
    paragraph = doc.add_paragraph()
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(paragraph.add_run(text), bold=True)
    keep_with_next(paragraph)


def add_lesson_heading(doc, lesson_number):
    heading = doc.add_heading(f"Bài {lesson_number}", level=2)
    keep_with_next(heading)
    for run in heading.runs:
        apply_run_font(run, bold=True)


def add_kanji_table(doc, items, mode, answer=False):
    rows = [mode["row"](item, answer=answer) for item in items]
    add_table(
        doc,
        mode["headers"],
        rows,
        column_widths=mode["widths"],
        table_alignment=WD_TABLE_ALIGNMENT.CENTER,
    )


def add_grouped_kanji(doc, items, mode):
    add_title(doc, "BÀI KIỂM TRA KANJI")
    add_subtitle(doc, mode["title"])

    current_lesson = None
    for lesson_number in sorted({item["lesson"] for item in items}):
        lesson_items = [item for item in items if item["lesson"] == lesson_number]
        if current_lesson is not None:
            doc.add_page_break()
        current_lesson = lesson_number
        add_lesson_heading(doc, lesson_number)
        add_kanji_table(doc, lesson_items, mode)


def add_grouped_kanji_with_answers(doc, items, mode):
    add_grouped_kanji(doc, items, mode)
    doc.add_page_break()
    add_title(doc, "ĐÁP ÁN")
    for index, lesson_number in enumerate(sorted({item["lesson"] for item in items})):
        if index > 0:
            doc.add_page_break()
        lesson_items = [item for item in items if item["lesson"] == lesson_number]
        add_lesson_heading(doc, lesson_number)
        add_kanji_table(doc, lesson_items, mode, answer=True)


def add_mixed_kanji(doc, items, mode):
    add_title(doc, "BÀI KIỂM TRA KANJI LỘN XỘN")
    add_subtitle(doc, mode["title"])
    add_kanji_table(doc, items, mode)


def add_mixed_kanji_with_answers(doc, items, mode):
    add_mixed_kanji(doc, items, mode)
    doc.add_page_break()
    add_title(doc, "ĐÁP ÁN")
    answer_headers = ["STT", "Bài", "Kanji", "Chữ mềm", "Nghĩa"]
    rows = [
        [str(index), str(item["lesson"]), item["kanji"], item["kana"], item["meaning"]]
        for index, item in enumerate(items, start=1)
    ]
    add_table(
        doc,
        answer_headers,
        rows,
        column_widths=MIXED_ANSWER_WIDTHS,
        table_alignment=WD_TABLE_ALIGNMENT.CENTER,
    )


def get_level_name(levels):
    if levels == {"N5"}:
        return "n5"
    if levels == {"N4"}:
        return "n4"
    return "n5_n4"


def get_lesson_name(lessons):
    lesson_values = sorted(lessons)
    if not lesson_values:
        return "all"
    if len(lesson_values) == 1:
        return f"bai{lesson_values[0]}"
    if lesson_values == list(range(lesson_values[0], lesson_values[-1] + 1)):
        return f"bai{lesson_values[0]}-{lesson_values[-1]}"
    return "bai" + "-".join(str(value) for value in lesson_values)


def get_available_output_path(output_dir, stem, suffix):
    path = output_dir / f"{stem}{suffix}"
    if not path.exists():
        return path

    counter = 2
    while True:
        path = output_dir / f"{stem}_{counter}{suffix}"
        if not path.exists():
            return path
        counter += 1


def build_output_stem(levels, mode, build_type, lessons):
    level_part = get_level_name(levels)
    type_names = {
        "1": "theo_bai_dap_an",
        "2": "tron_dap_an",
        "3": "theo_bai",
        "4": "tron",
    }
    type_part = type_names[build_type]
    lesson_part = get_lesson_name(lessons)
    return f"kanji_{level_part}_{lesson_part}_{type_part}_{mode['name']}"


def build_kanji_document():
    levels = ask_kanji_levels()
    mode = ask_kanji_mode()
    build_type = ask_kanji_build_type()
    items = get_kanji_items(levels)

    if not items:
        print("Không tìm thấy dữ liệu Kanji phù hợp với cấp độ đã chọn.")
        return

    available_lessons = sorted({item["lesson"] for item in items})
    lessons = ask_lessons(available_lessons)
    selected_items = [item for item in items if item["lesson"] in lessons]

    if build_type in {"2", "4"}:
        random.shuffle(selected_items)
        selected_items = selected_items[:ask_mixed_count(len(selected_items))]

    doc = setup_document(footer_text="Bài kiểm tra Kanji", booklet_layout=False)
    if build_type == "1":
        add_grouped_kanji_with_answers(doc, selected_items, mode)
    elif build_type == "2":
        add_mixed_kanji_with_answers(doc, selected_items, mode)
    elif build_type == "3":
        add_grouped_kanji(doc, selected_items, mode)
    else:
        add_mixed_kanji(doc, selected_items, mode)

    docx_dir = OUTPUT_DIR / "docx" / "kanji"
    pdf_dir = OUTPUT_DIR / "pdf" / "kanji"
    docx_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir.mkdir(parents=True, exist_ok=True)

    output_stem = build_output_stem(levels, mode, build_type, lessons)
    docx_path = get_available_output_path(docx_dir, output_stem, ".docx")
    pdf_path = pdf_dir / f"{docx_path.stem}.pdf"

    doc.save(docx_path)
    print(f"Đã tạo file Word Kanji: {docx_path}")
    update_docx_fields_and_export_pdf(docx_path, pdf_path)
    print(f"Đã tạo file PDF Kanji: {pdf_path}")

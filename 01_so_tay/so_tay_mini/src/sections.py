from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from docx_utils import (
    add_simple_field,
    add_table,
    apply_paragraph_format,
    apply_run_font,
    keep_with_next,
)


INDEX_EMPTY_RIGHT_PLACEHOLDER = "\u00a0"
INDEX_LINE_WEIGHT_LIMIT = 55
INDEX_LONG_LEFT_WEIGHT_LIMIT = 45
INDEX_SHORT_RIGHT_WEIGHT_LIMIT = 12
INDEX_LONG_RIGHT_WEIGHT_LIMIT = 28

ROMAN_NUMERALS = (
    "I",
    "II",
    "III",
    "IV",
    "V",
    "VI",
    "VII",
    "VIII",
    "IX",
    "X",
)


def roman_label(index):
    if 1 <= index <= len(ROMAN_NUMERALS):
        return ROMAN_NUMERALS[index - 1]
    return str(index)


def strip_number_prefix(text):
    return text.split(".", 1)[1].strip() if text and text.split(".", 1)[0].isdigit() else text


def strip_slash_suffix(text):
    return text.split("/", 1)[0].strip() if text else text


def add_roman_heading(doc, index, text, level=2):
    clean_text = strip_number_prefix(text or "Nội dung")
    heading = doc.add_heading(f"{roman_label(index)}. {clean_text}", level=level)
    keep_with_next(heading)
    for run in heading.runs:
        apply_run_font(run, bold=True)
    return heading


def add_toc_section(doc):
    title = doc.add_paragraph()
    apply_paragraph_format(title, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(title.add_run("MỤC LỤC"), bold=True)

    toc_paragraph = doc.add_paragraph()
    apply_paragraph_format(toc_paragraph)
    add_simple_field(toc_paragraph, r'TOC \o "1-3" \h \z \u')


def add_section_with_break(doc, add_func, *args):
    if len(doc.paragraphs) > 0 or len(doc.tables) > 0:
        doc.add_page_break()
    add_func(doc, *args)


def add_5s_paragraphs(doc, items):
    for index, item in enumerate(items):
        if index > 0:
            doc.add_paragraph()

        jp_paragraph = doc.add_paragraph(style="Bang_Style")
        apply_paragraph_format(jp_paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY)
        apply_run_font(jp_paragraph.add_run(item.get("tu_vung", "")))
        keep_with_next(jp_paragraph)

        text = item.get("y_nghia", "")
        if ":" in text:
            prefix, suffix = text.split(":", 1)
            prefix_paragraph = doc.add_paragraph(style="Bang_Style")
            apply_paragraph_format(prefix_paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY)
            apply_run_font(prefix_paragraph.add_run(f"{prefix}: "), bold=True)
            if suffix.strip():
                keep_with_next(prefix_paragraph)
                detail_paragraph = doc.add_paragraph(style="Bang_Style")
                apply_paragraph_format(detail_paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY)
                apply_run_font(detail_paragraph.add_run(suffix.lstrip()))
        else:
            vn_paragraph = doc.add_paragraph(style="Bang_Style")
            apply_paragraph_format(vn_paragraph, WD_ALIGN_PARAGRAPH.JUSTIFY)
            apply_run_font(vn_paragraph.add_run(text))


def add_text_paragraph(doc, text, justify=True):
    paragraph = doc.add_paragraph(style="Bang_Style")
    alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    apply_paragraph_format(paragraph, alignment)
    apply_run_font(paragraph.add_run(text))
    return paragraph


def add_prefix_bold_runs(paragraph, text):
    if ":" not in text:
        apply_run_font(paragraph.add_run(text))
        return

    prefix, suffix = text.split(":", 1)
    apply_run_font(paragraph.add_run(f"{prefix}: "), bold=True)
    apply_run_font(paragraph.add_run(suffix.lstrip()))


def add_prefix_bold_paragraph(doc, text, justify=True):
    paragraph = doc.add_paragraph(style="Bang_Style")
    alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    apply_paragraph_format(paragraph, alignment)
    add_prefix_bold_runs(paragraph, text)
    return paragraph


def get_vocabulary_text(item):
    return item.get("tu_vung_hien_thi") or item.get("tu_vung", "")


def get_vietnamese_index_vocabulary_text(item):
    vocabulary_text = get_vocabulary_text(item)
    level_suffix = f" - {item.get('cap_do', '')}"
    if item.get("cap_do") and vocabulary_text.endswith(level_suffix):
        return vocabulary_text[: -len(level_suffix)]
    return vocabulary_text


def get_vietnamese_index_meaning_text(item):
    meaning_text = item.get("y_nghia", "")
    if item.get("cap_do") == "N4":
        return f"{meaning_text} - {item.get('cap_do')}"
    return meaning_text


def get_text_weight(text):
    weight = 0
    for char in text or "":
        if "\u3040" <= char <= "\u30ff" or "\u3400" <= char <= "\u9fff":
            weight += 2
        elif ord(char) > 127:
            weight += 1.2
        else:
            weight += 1
    return weight


def split_index_meaning_if_long(left_text, right_text, keep_medium_right_with_long_left=False):
    line_weight = get_text_weight(left_text) + get_text_weight(right_text)
    if line_weight <= INDEX_LINE_WEIGHT_LIMIT:
        return right_text, ""

    right_weight = get_text_weight(right_text)
    if (
        keep_medium_right_with_long_left
        and
        INDEX_SHORT_RIGHT_WEIGHT_LIMIT < right_weight < INDEX_LONG_RIGHT_WEIGHT_LIMIT
        and get_text_weight(left_text) >= INDEX_LONG_LEFT_WEIGHT_LIMIT
    ):
        return right_text, ""

    return "", right_text


def add_index_paragraph(
    doc,
    left_text,
    right_text,
    split_long_meaning=False,
    keep_medium_right_with_long_left=False,
):
    display_right = right_text
    continuation = ""
    if split_long_meaning:
        display_right, continuation = split_index_meaning_if_long(
            left_text,
            right_text,
            keep_medium_right_with_long_left=keep_medium_right_with_long_left,
        )

    paragraph = doc.add_paragraph(style="Index_Style")
    apply_paragraph_format(paragraph)
    if continuation and not display_right:
        display_right = INDEX_EMPTY_RIGHT_PLACEHOLDER
    paragraph.text = f"{left_text}\t{display_right}"

    if continuation:
        paragraph.paragraph_format.keep_with_next = True
        continuation_paragraph = doc.add_paragraph(style="Index_Style")
        apply_paragraph_format(continuation_paragraph)
        continuation_paragraph.text = f"\t{continuation}"


def add_text_lines(doc, text_or_lines, justify=True):
    paragraphs = []
    if isinstance(text_or_lines, list):
        for text in text_or_lines:
            paragraphs.append(add_text_paragraph(doc, text, justify=justify))
    elif text_or_lines:
        paragraphs.append(add_text_paragraph(doc, text_or_lines, justify=justify))
    return paragraphs


def add_text_pair(doc, japanese_text, vietnamese_text):
    japanese_paragraphs = add_text_lines(doc, japanese_text)
    vietnamese_paragraph = None
    if vietnamese_text:
        vietnamese_paragraph = add_text_paragraph(doc, vietnamese_text)
        for paragraph in japanese_paragraphs:
            keep_with_next(paragraph)
    return japanese_paragraphs, vietnamese_paragraph


def add_text_pair_with_bold_translation_prefix(doc, japanese_text, vietnamese_text):
    japanese_paragraphs = add_text_lines(doc, japanese_text)
    vietnamese_paragraph = None
    if vietnamese_text:
        vietnamese_paragraph = add_prefix_bold_paragraph(doc, vietnamese_text)
        for paragraph in japanese_paragraphs:
            keep_with_next(paragraph)
    return japanese_paragraphs, vietnamese_paragraph


def add_inline_title(doc, text):
    paragraph = doc.add_paragraph(style="Bang_Style")
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.LEFT)
    apply_run_font(paragraph.add_run(text), bold=True)
    keep_with_next(paragraph)
    return paragraph


def add_centered_text(doc, text):
    paragraph = doc.add_paragraph(style="Bang_Style")
    apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(paragraph.add_run(text))
    keep_with_next(paragraph)
    return paragraph


def add_online_version_page(doc, qr_path, updated_date):
    def add_centered_line(text, bold=False):
        paragraph = doc.add_paragraph()
        apply_paragraph_format(paragraph, WD_ALIGN_PARAGRAPH.CENTER)
        apply_run_font(paragraph.add_run(text), bold=bold)
        return paragraph

    title = doc.add_paragraph()
    apply_paragraph_format(title, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(title.add_run("Phiên bản Online của tài liệu:"), bold=True)

    qr_paragraph = doc.add_paragraph()
    apply_paragraph_format(qr_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    if qr_path.exists():
        qr_paragraph.add_run().add_picture(str(qr_path), width=Cm(5), height=Cm(5))
    else:
        apply_run_font(qr_paragraph.add_run("[Thiếu ảnh QR]"), bold=True)

    date_paragraph = doc.add_paragraph()
    apply_paragraph_format(date_paragraph, WD_ALIGN_PARAGRAPH.CENTER)
    apply_run_font(date_paragraph.add_run(f"Cập nhật lần cuối: {updated_date}"))

    add_centered_line("著者（ちょしゃ）：", bold=True)
    add_centered_line("ズオン・タイン・タン")
    add_centered_line("編集協力（へんしゅうきょうりょく）：", bold=True)
    add_centered_line("ChatGPT 5.5")


def add_horenso_section(doc, data):
    if not data:
        return

    keep_with_next(doc.add_heading("HORENSO", level=1))

    for index, item in enumerate(data.get("noi_dung", []), start=1):
        tu_vung = item.get("tu_vung_tieng_nhat", "")
        nghia = item.get("nghia_tieng_viet", "")
        add_roman_heading(doc, index, nghia.capitalize())
        add_text_pair_with_bold_translation_prefix(
            doc,
            item.get("giai_thich_tieng_nhat", ""),
            f"{tu_vung}: {item.get('giai_thich_tieng_viet', '')}",
        )


def add_aisatsu_section(doc, data):
    if not data:
        return

    keep_with_next(doc.add_heading("AISATSU", level=1))
    for index, section in enumerate(data, start=1):
        title_vn = section.get("tieu_de_tieng_viet", "")
        heading = add_roman_heading(doc, index, title_vn)
        if index > 1 and section.get("ngat_trang_truoc", False):
            heading.paragraph_format.page_break_before = True

        for item in section.get("noi_dung", []):
            if item.get("tieu_de"):
                add_inline_title(doc, item.get("tieu_de", ""))
            add_text_pair(doc, item.get("jp", ""), item.get("vi", ""))


def add_garbage_section(doc, data):
    if not data:
        return

    keep_with_next(doc.add_heading("PHÂN LOẠI RÁC", level=1))
    for index, section in enumerate(data, start=1):
        title_vn = section.get("tieu_de_tieng_viet", "")
        title_jp = strip_slash_suffix(section.get("tieu_de_tieng_nhat", ""))
        heading = add_roman_heading(doc, index, title_vn)
        if index > 1:
            heading.paragraph_format.page_break_before = True

        add_centered_text(doc, title_jp)
        rows = [
            (item.get("cau_tieng_nhat", ""), item.get("cau_tieng_viet", ""))
            for item in section.get("noi_dung", [])
        ]
        add_table(doc, ["Tiếng Nhật", "Tiếng Việt"], rows)


def add_japanese_index(doc, all_words):
    keep_with_next(doc.add_heading("TRA NHANH (HIRA, KATA)", level=1))
    current_group = ""
    for item in sorted(all_words, key=lambda row: row["jp_sort_key"]):
        group_char = item["jp_group_char"]
        if group_char != current_group:
            current_group = group_char
            group_paragraph = doc.add_paragraph()
            apply_paragraph_format(group_paragraph)
            apply_run_font(group_paragraph.add_run(current_group), bold=True)
            keep_with_next(group_paragraph)

        add_index_paragraph(
            doc,
            get_vocabulary_text(item),
            item.get("y_nghia", ""),
            split_long_meaning=True,
            keep_medium_right_with_long_left=True,
        )


def add_vietnamese_index(doc, all_words):
    keep_with_next(doc.add_heading("TRA NHANH (TIẾNG VIỆT)", level=1))
    current_group = ""
    for item in sorted(all_words, key=lambda row: row["vn_sort_key"]):
        group_char = item["vn_group_char"]
        if group_char != current_group:
            current_group = group_char
            group_paragraph = doc.add_paragraph()
            apply_paragraph_format(group_paragraph)
            apply_run_font(group_paragraph.add_run(current_group), bold=True)
            keep_with_next(group_paragraph)

        add_index_paragraph(
            doc,
            get_vietnamese_index_meaning_text(item),
            get_vietnamese_index_vocabulary_text(item),
            split_long_meaning=True,
        )


def add_5s_section(doc, data):
    if not data:
        return

    keep_with_next(doc.add_heading("NGUYÊN TẮC 5S-K-C", level=1))
    grouped = {}
    for item in data:
        grouped.setdefault(item.get("phan_nhom", "Khác"), []).append(item)

    for index, (group_name, items) in enumerate(grouped.items(), start=1):
        heading = add_roman_heading(doc, index, f"Nguyên tắc {group_name}")
        if index > 1:
            heading.paragraph_format.page_break_before = True

        if group_name == "5S":
            add_5s_paragraphs(doc, items)
        else:
            rows = [(item.get("tu_vung", ""), item.get("y_nghia", "")) for item in items]
            add_table(doc, None, rows, column_widths=[Cm(6.19), Cm(5.49)], border_color="FFFFFF")


def add_grouped_vocabulary(doc, json_data_map):
    keep_with_next(doc.add_heading("TỪ VỰNG N5 - N4", level=1))
    for index, file_data in enumerate(json_data_map):
        heading_text = file_data["display_title"]
        heading_index = file_data["prefix_num"]

        heading = doc.add_heading(heading_text, level=2)
        keep_with_next(heading)
        if index > 0:
            heading.paragraph_format.page_break_before = True

        grouped_words = {}
        for item in file_data["data"]:
            group = item.get("phan_nhom", "Khác")
            grouped_words.setdefault(group, []).append(item)

        sub_index = 1
        for group, items in grouped_words.items():
            sub_heading = doc.add_heading(f"{heading_index}.{sub_index}. {group}", level=3)
            keep_with_next(sub_heading)
            if sub_index > 1:
                sub_heading.paragraph_format.page_break_before = True

            sub_index += 1
            rows = [(get_vocabulary_text(item), item.get("y_nghia", "")) for item in items]
            add_table(doc, ["Từ vựng", "Ý nghĩa"], rows)

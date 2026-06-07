import sys
import unittest
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT_DIR = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT_DIR / "src"
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from data_loader import load_kaiwa_before_departure
from kaiwa_builder import KAIWA_DOCX_NAME, KAIWA_PDF_NAME, create_kaiwa_document


def contains_kanji(text):
    return any("\u3400" <= char <= "\u9fff" for char in text or "")


class KaiwaBeforeDepartureTest(unittest.TestCase):
    def test_json_has_expected_structure(self):
        data = load_kaiwa_before_departure(ROOT_DIR / "data")

        self.assertEqual(data["title"], "KAIWA TRƯỚC XUẤT CẢNH")
        self.assertEqual(len(data["sections"]), 7)
        self.assertEqual(sum(len(section["items"]) for section in data["sections"]), 22)

    def test_japanese_qa_text_has_no_kanji(self):
        data = load_kaiwa_before_departure(ROOT_DIR / "data")

        for section in data["sections"]:
            for item in section["items"]:
                for qa in item["qa"]:
                    for key in ["question_jp", "answer_jp"]:
                        with self.subTest(section=section["title_vi"], item=item["title_vi"], key=key):
                            self.assertFalse(contains_kanji(qa[key]), qa[key])

    def test_suru_verbs_are_not_split_after_sino_japanese_nouns(self):
        data = load_kaiwa_before_departure(ROOT_DIR / "data")
        forbidden_phrases = [
            "せいとん します",
            "せいそう します",
            "そうじ します",
            "ほうこく します",
            "れんらく します",
            "そうだん します",
            "そうだん です",
        ]

        for section in data["sections"]:
            for item in section["items"]:
                for qa in item["qa"]:
                    for key in ["question_jp", "answer_jp"]:
                        for phrase in forbidden_phrases:
                            with self.subTest(
                                section=section["title_vi"],
                                item=item["title_vi"],
                                key=key,
                                phrase=phrase,
                            ):
                                self.assertNotIn(phrase, qa[key])

    def test_document_is_a4_and_uses_no_tables(self):
        data = load_kaiwa_before_departure(ROOT_DIR / "data")
        doc = create_kaiwa_document(data)
        section = doc.sections[0]
        paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]

        self.assertEqual(len(doc.tables), 0)
        self.assertAlmostEqual(section.page_width.cm, 21.0, places=1)
        self.assertAlmostEqual(section.page_height.cm, 29.7, places=1)
        self.assertIn("KAIWA TRƯỚC XUẤT CẢNH", paragraph_texts)
        self.assertTrue(any("じこしょうかいを して ください。" in paragraph.text for paragraph in doc.paragraphs))

    def test_document_does_not_show_redundant_qa_labels(self):
        data = load_kaiwa_before_departure(ROOT_DIR / "data")
        doc = create_kaiwa_document(data)
        paragraph_texts = [paragraph.text for paragraph in doc.paragraphs]

        self.assertIn("じこしょうかいを して ください。", paragraph_texts)
        self.assertIn("Hãy giới thiệu bản thân.", paragraph_texts)
        self.assertFalse(any(text.startswith("Q1.") for text in paragraph_texts))
        self.assertFalse(any(text.startswith("A1.") for text in paragraph_texts))
        self.assertNotIn("Nghĩa: Hãy giới thiệu bản thân.", paragraph_texts)

    def test_output_file_names_include_docx_and_pdf(self):
        self.assertEqual(KAIWA_DOCX_NAME, "kaiwa_truoc_xuat_canh.docx")
        self.assertEqual(KAIWA_PDF_NAME, "kaiwa_truoc_xuat_canh.pdf")

    def test_intro_description_is_centered(self):
        data = load_kaiwa_before_departure(ROOT_DIR / "data")
        doc = create_kaiwa_document(data)
        description = data["description"]
        paragraph = next(paragraph for paragraph in doc.paragraphs if paragraph.text == description)

        self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)


if __name__ == "__main__":
    unittest.main()

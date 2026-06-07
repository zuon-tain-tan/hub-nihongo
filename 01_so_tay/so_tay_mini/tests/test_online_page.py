import sys
import unittest
from pathlib import Path

from docx import Document


ROOT_DIR = Path(__file__).resolve().parents[1]
SRC_DIR = ROOT_DIR / "src"

if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

from sections import add_completion_image_page, add_online_version_page


class OnlineVersionPageTest(unittest.TestCase):
    def test_author_and_editor_are_split_without_spacers(self):
        doc = Document()

        add_online_version_page(doc, Path("missing-qr.png"), "07/06/2026")

        texts = [paragraph.text for paragraph in doc.paragraphs]
        self.assertEqual(
            texts[-5:],
            [
                "Cập nhật lần cuối: 07/06/2026",
                "著者（ちょしゃ）：",
                "ズオン・タイン・タン",
                "編集協力（へんしゅうきょうりょく）：",
                "ChatGPT 5.5",
            ],
        )
        self.assertNotIn("", texts[-5:])

    def test_completion_image_page_uses_18_cm_height(self):
        doc = Document()
        image_path = ROOT_DIR / "assets" / "images" / "notebook" / "hoan_thanh_so_tay.png"

        add_completion_image_page(doc, image_path)

        self.assertEqual(len(doc.inline_shapes), 1)
        self.assertAlmostEqual(doc.inline_shapes[0].height.cm, 18.0, places=1)


if __name__ == "__main__":
    unittest.main()
